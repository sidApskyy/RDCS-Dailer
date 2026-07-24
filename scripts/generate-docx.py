#!/usr/bin/env python3
"""
Generate a merged Microsoft Word (.docx) version of the RDCS In-House Dialer
Platform architecture package.

Requires:
  - pandoc (https://pandoc.org/installing.html)
  - pypandoc (pip install pypandoc)
  - python-docx (pip install python-docx)

Usage:
  python scripts/generate-docx.py

Output:
  RDCS-In-House-Dialer-Platform.docx in the project root.
"""

import os
import subprocess
import sys
from datetime import datetime
from pathlib import Path


def ensure_dependencies():
    try:
        import pypandoc
        from docx import Document
    except ImportError as e:
        print(f"Missing dependency: {e}")
        print("Install required packages with:")
        print("  pip install -r scripts/requirements.txt")
        sys.exit(1)


def get_docs(root: Path) -> list[Path]:
    docs_dir = root / "docs"
    files = sorted(docs_dir.glob("*.md"))
    # Filter to numbered files 00-73 to preserve order and exclude stray files
    numbered = []
    for f in files:
        name = f.stem
        if name[:2].isdigit() and name[2] == "-":
            numbered.append(f)
    numbered.sort(key=lambda p: int(p.stem[:2]))
    return numbered


def build_front_matter(root: Path) -> str:
    today = datetime.now().strftime("%d-%b-%Y")
    return f"""---
title: "RDCS In-House Dialer Platform"
subtitle: "Enterprise Architecture & Engineering Implementation Package"
author: "Enterprise Architecture Team"
date: "{today}"
version: "1.0.0"
status: "Draft"
---

# Cover Page

**RDCS In-House Dialer Platform**

**Enterprise Architecture & Engineering Implementation Package**

Prepared by: Enterprise Architecture Team

Version: 1.0.0

Status: Draft

Date: {today}

Classification: Internal – Engineering Use

---

# Document Control

| Property | Value |
|----------|-------|
| Project Name | RDCS In-House Dialer Platform |
| Document Title | Enterprise Architecture & Engineering Implementation Package |
| Version | 1.0.0 |
| Status | Draft |
| Classification | Internal – Engineering Use |
| Last Updated | {today} |
| Document Owner | Enterprise Architecture Team |
| Review Cycle | Quarterly |

---

# Revision History

| Version | Date | Author | Description |
|---------|------|--------|-------------|
| 1.0.0 | {today} | Enterprise Architecture Team | Initial release of complete architecture package |

---

# Table of Contents

"""


def build_appendix() -> str:
    return """

---

# Appendices

## Appendix A: Glossary and Acronyms

See `72-glossary-acronyms.md` for the complete glossary and acronym list.

## Appendix B: Permission Matrix

See `09-permission-matrix.md` for the complete role and permission matrix.

## Appendix C: Database Schema

See `36-prisma-schema-design.md` for the full Prisma schema.

## Appendix D: API Endpoint Reference

See `40-rest-api-documentation.md` for all REST endpoints and `41-websocket-api-documentation.md` for WebSocket events.

## Appendix E: Event Catalog

See `42-internal-event-documentation.md` and `43-webhook-events.md` for the event catalog.

## Appendix F: Production Checklist

See `60-production-checklist.md` for the go-live readiness checklist.

## Appendix G: References

See `73-references.md` for external references, standards, and tools.

"""


def main():
    ensure_dependencies()

    import pypandoc
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    root = Path(__file__).resolve().parent.parent
    docs = get_docs(root)

    if not docs:
        print(f"No docs found in {root / 'docs'}")
        sys.exit(1)

    print(f"Found {len(docs)} documentation files.")

    # Build a single markdown string
    combined = build_front_matter(root)

    for doc in docs:
        print(f"  - {doc.name}")
        combined += f"\n\n\n---\n\n"
        with open(doc, "r", encoding="utf-8") as f:
            combined += f.read()

    combined += build_appendix()

    output_path = root / "RDCS-In-House-Dialer-Platform.docx"

    combined_md_path = root / "_combined.md"
    combined_md_path.write_text(combined, encoding="utf-8")
    print(f"\nWrote combined markdown to: {combined_md_path}")

    print(f"Generating DOCX with pandoc: {output_path}")
    try:
        result = subprocess.run(
            [
                "pandoc",
                str(combined_md_path),
                "-f",
                "markdown",
                "-t",
                "docx",
                "--toc",
                "--toc-depth=3",
                "-o",
                str(output_path),
            ],
            capture_output=True,
            text=True,
            check=True,
        )
        if result.stderr:
            print(result.stderr)
    except subprocess.CalledProcessError as e:
        print(f"pandoc conversion failed with exit code {e.returncode}")
        print("STDOUT:", e.stdout)
        print("STDERR:", e.stderr)
        sys.exit(1)
    finally:
        combined_md_path.unlink(missing_ok=True)
        print(f"Removed temporary markdown file: {combined_md_path}")

    # Post-process with python-docx: add cover page header/footer, ensure page break after cover
    try:
        doc = Document(str(output_path))

        # Add a header to all sections
        for section in doc.sections:
            header = section.header
            header_para = header.paragraphs[0]
            header_para.text = "RDCS In-House Dialer Platform | Confidential"
            header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in header_para.runs:
                run.font.size = Pt(8)
                run.font.italic = True

            footer = section.footer
            footer_para = footer.paragraphs[0]
            footer_para.text = "Document Version 1.0.0 | Page "
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in footer_para.runs:
                run.font.size = Pt(8)
            # Add page number field
            footer_para.add_run().add_field("PAGE")
            footer_para.add_run(" of ")
            footer_para.add_run().add_field("NUMPAGES")

        doc.save(str(output_path))
        print("Added headers and footers.")
    except Exception as e:
        print(f"Warning: could not add headers/footers: {e}")

    print(f"\nSuccessfully generated: {output_path}")


if __name__ == "__main__":
    main()
